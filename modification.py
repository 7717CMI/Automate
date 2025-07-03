from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, jsonify, session
import pandas as pd
import os
import subprocess
import time
import pyautogui
from functools import wraps
from werkzeug.utils import secure_filename
import win32com.client
from docx import Document
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager
import threading
import logging
import re
import requests
import shutil

app = Flask(__name__)
app.secret_key = 'your_secret_key_change_this_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Hardcoded download filename and path
app.config['DOWNLOAD_FILENAME'] = 'ROB.xlsx'
app.config['DOWNLOAD_PATH'] = r'C:\Users\akshat\Desktop\RPA\\' + app.config['DOWNLOAD_FILENAME']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}

# Set up logging to capture output
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a global variable to track processing status
processing_status = {
    'active': False,
    'message': 'Ready',
    'progress': 0,
    'total': 0,
    'current_file': '',
    'logs': []
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def log_to_status(message):
    """Add a message to the processing status logs"""
    global processing_status
    processing_status['logs'].append(f"{datetime.now().strftime('%H:%M:%S')}: {message}")
    print(f"[LOG] {message}")  # Also print to console

# NEW: Google Trends Agent Class
class GoogleTrendsAgent:
    def __init__(self, serp_api_key):
        self.api_key = serp_api_key
        self.base_url = "https://serpapi.com/search.json"

    def get_trends_data(self, keyword):
        params = {
            'engine': 'google_trends',
            'q': keyword,
            'data_type': 'TIMESERIES',
            'date': 'now 7-d',
            'geo': '',
            'api_key': self.api_key
        }
        try:
            response = requests.get(self.base_url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            if 'error' in data:
                log_to_status(f"[API ERROR] Google Trends error for '{keyword}': {data['error']}")
                return []
            return data.get('interest_over_time', {}).get('timeline_data', [])
        except Exception as e:
            log_to_status(f"[API ERROR] Exception for '{keyword}': {e}")
            return []

    def get_region_data(self, keyword):
        params = {
            'engine': 'google_trends',
            'q': keyword,
            'data_type': 'GEO_MAP_0',
            'date': 'now 7-d',
            'api_key': self.api_key
        }
        try:
            response = requests.get(self.base_url, params=params, timeout=30)
            response.raise_for_status()
            data = response.json()
            if 'error' in data:
                log_to_status(f"[API ERROR] Google Trends region error for '{keyword}': {data['error']}")
                return []
            return data.get('interest_by_region', [])
        except Exception as e:
            log_to_status(f"[API ERROR] Exception (region) for '{keyword}': {e}")
            return []

    def check_threshold_criteria(self, interest_values):
        if not interest_values:
            return False
        days_above_50 = sum(1 for value in interest_values if value and value > 50)
        return days_above_50 >= 2

    def check_country_criteria(self, region_data):
        target_countries = ['United States', 'Germany', 'Japan', 'United Kingdom', 'France', 'China']
        top_countries = [item.get('location', '') for item in region_data[:10]]
        return any(country in top_countries for country in target_countries)

    def extract_interest_values(self, trends_data, keyword):
        interest_values = []
        for item in trends_data:
            values = item.get('values', [])
            if values:
                if len(values) == 1:
                    interest_values.append(values[0].get('extracted_value', 0))
                else:
                    for value in values:
                        if value.get('query', '').lower() == keyword.lower():
                            interest_values.append(value.get('extracted_value', 0))
                            break
                    else:
                        interest_values.append(0)
        return interest_values

    def analyze_keyword(self, keyword):
        trends_data = self.get_trends_data(keyword)
        if not trends_data:
            return None
        interest_values = self.extract_interest_values(trends_data, keyword)
        if len(interest_values) < 7:
            return None
        time.sleep(1)
        region_data = self.get_region_data(keyword)
        meets_threshold = self.check_threshold_criteria(interest_values)
        has_target_countries = self.check_country_criteria(region_data)
        if meets_threshold and has_target_countries:
            return {
                'keyword': keyword,
                'interest_values': interest_values,
                'max_interest': max(interest_values) if interest_values else 0,
                'avg_interest': round(sum(interest_values) / len(interest_values)) if interest_values else 0,
                'days_above_50': sum(1 for v in interest_values if v and v > 50),
                'top_countries': [item.get('location', '') for item in region_data[:5]]
            }
        return None

# Home route
@app.route('/')
def index():
    return render_template('index.html')

# NEW: Weekly Keyword Analysis Routes
@app.route('/weekly_keyword_analysis', methods=['GET', 'POST'])
def weekly_keyword_analysis():
    """Weekly keyword analysis with ranking sheet and backend ROB data processing"""
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data
            api_key = request.form.get('api_key')
            ranking_file = request.files.get('ranking_file')
            backend_rob_file = request.files.get('backend_rob_file')
            
            # Validate inputs
            if not api_key:
                flash('SERP API key is required!')
                return redirect(request.url)
            
            if not ranking_file or ranking_file.filename == '':
                flash('Ranking sheet file is required!')
                return redirect(request.url)
                
            if not backend_rob_file or backend_rob_file.filename == '':
                flash('Backend ROB data file is required!')
                return redirect(request.url)
            
            # Validate file types
            if not allowed_file(ranking_file.filename) or not allowed_file(backend_rob_file.filename):
                flash('Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)
            
            # Save uploaded files
            ranking_filename = secure_filename(f"ranking_{ranking_file.filename}")
            backend_filename = secure_filename(f"backend_{backend_rob_file.filename}")
            
            ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
            backend_path = os.path.join(app.config['UPLOAD_FOLDER'], backend_filename)
            
            ranking_file.save(ranking_path)
            backend_rob_file.save(backend_path)
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting weekly keyword analysis...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread
            threading.Thread(target=process_weekly_keywords_with_feedback, 
                           args=(api_key, ranking_path, backend_path)).start()
            
            flash('Weekly keyword analysis started! Check the status page for updates.')
            return redirect(url_for('processing_status_page'))
            
        except Exception as e:
            flash(f'Error starting keyword analysis: {str(e)}')
            logger.error(f"Error in weekly_keyword_analysis: {e}")
            return render_template('weekly_keyword_analysis.html')
    
    return render_template('weekly_keyword_analysis.html')

@app.route('/download_weekly_results/<filename>')
def download_weekly_results(filename):
    """Download weekly analysis result files"""
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            flash(f'File not found: {filename}')
            return redirect(url_for('weekly_keyword_analysis'))
    except Exception as e:
        flash(f'Error downloading file: {str(e)}')
        return redirect(url_for('weekly_keyword_analysis'))

@app.route('/weekly_analysis_history')
def weekly_analysis_history():
    """Show history of weekly keyword analyses"""
    try:
        upload_folder = app.config['UPLOAD_FOLDER']
        files = []
        
        for filename in os.listdir(upload_folder):
            if filename.startswith(('qualified_keywords_', 'qualified_rob_data_', 'updated_ranking_sheet_')):
                file_path = os.path.join(upload_folder, filename)
                file_info = {
                    'name': filename,
                    'size': os.path.getsize(file_path),
                    'created': datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
                }
                files.append(file_info)
        
        # Sort by creation date (newest first)
        files.sort(key=lambda x: x['created'], reverse=True)
        
        return render_template('weekly_analysis_history.html', files=files)
        
    except Exception as e:
        flash(f'Error loading history: {str(e)}')
        return redirect(url_for('weekly_keyword_analysis'))

# NEW: Weekly Keyword Processing Function
def process_weekly_keywords_with_feedback(api_key, ranking_path, backend_path):
    """Process weekly keywords with status feedback"""
    global processing_status
    
    try:
        log_to_status("Starting weekly keyword analysis...")
        
        # Step 1: Load ranking sheet
        log_to_status("Loading ranking sheet...")
        ranking_df = pd.read_excel(ranking_path)
        
        # Validate required columns in ranking sheet
        required_ranking_columns = ['Keywords', 'AVG. Search', 'Competition', 'RID']
        missing_columns = [col for col in required_ranking_columns if col not in ranking_df.columns]
        
        if missing_columns:
            processing_status['active'] = False
            processing_status['message'] = f"Missing columns in ranking sheet: {', '.join(missing_columns)}"
            log_to_status(f"ERROR: Missing columns in ranking sheet: {', '.join(missing_columns)}")
            return
        
        # Step 2: Load backend ROB data
        log_to_status("Loading backend ROB data...")
        backend_df = pd.read_excel(backend_path)
        
        # Step 3: Filter keywords based on criteria
        log_to_status("Filtering keywords based on search volume and competition...")
        ranking_df['AVG. Search'] = pd.to_numeric(ranking_df['AVG. Search'], errors='coerce')
        filtered_df = ranking_df[(ranking_df['AVG. Search'] >= 5000) & (ranking_df['Competition'] == 'Low')]
        
        # Remove 'Market' from keywords
        filtered_df = filtered_df.copy()
        filtered_df['Keywords'] = filtered_df['Keywords'].str.replace(r'\s*Market$', '', regex=True)
        
        keywords = filtered_df['Keywords'].dropna().tolist()
        processing_status['total'] = len(keywords)
        log_to_status(f"Found {len(keywords)} keywords meeting basic criteria")
        
        # Step 4: Analyze keywords with Google Trends
        log_to_status("Starting Google Trends analysis...")
        agent = GoogleTrendsAgent(api_key)
        qualifying_keywords = []
        
        for i, keyword in enumerate(keywords):
            processing_status['progress'] = i
            processing_status['current_file'] = keyword
            processing_status['message'] = f"Analyzing keyword {i+1} of {len(keywords)}: {keyword}"
            try:
                result = agent.analyze_keyword(keyword)
                if result:
                    qualifying_keywords.append(result)
                    log_to_status(f"✅ Keyword qualified: {keyword}")
                else:
                    log_to_status(f"❌ Keyword rejected: {keyword} (Did not meet Google Trends criteria or insufficient data)")
            except Exception as e:
                log_to_status(f"[ERROR] Exception analyzing '{keyword}': {e}")
            # Check if we've reached the 300 limit
            if len(qualifying_keywords) >= 300:
                log_to_status("⚠️ Reached maximum limit of 300 qualified keywords")
                break
        
        if not qualifying_keywords:
            processing_status['active'] = False
            processing_status['message'] = "No keywords qualified for trends criteria"
            log_to_status("No keywords qualified for Google Trends criteria")
            return
        
        # Step 5: Get qualifying RIDs
        log_to_status("Extracting qualifying RIDs...")
        qualifying_keyword_list = [result['keyword'] for result in qualifying_keywords]
        qualifying_rid_df = filtered_df[filtered_df['Keywords'].isin(qualifying_keyword_list)]
        qualifying_rids = qualifying_rid_df['RID'].tolist()
        
        log_to_status(f"Found {len(qualifying_rids)} qualifying RIDs")
        
        # Step 6: Extract data from backend ROB for qualifying RIDs
        log_to_status("Extracting data from backend ROB for qualifying RIDs...")
        if 'RID' in backend_df.columns:
            qualified_rob_data = backend_df[backend_df['RID'].isin(qualifying_rids)]
        else:
            # Try alternative column names
            rid_column = None
            for col in backend_df.columns:
                if 'rid' in col.lower() or ('report' in col.lower() and 'id' in col.lower()):
                    rid_column = col
                    break
            
            if rid_column:
                qualified_rob_data = backend_df[backend_df[rid_column].isin(qualifying_rids)]
            else:
                processing_status['active'] = False
                processing_status['message'] = "Could not find RID column in backend ROB data"
                log_to_status("ERROR: Could not find RID column in backend ROB data")
                return
        
        # Step 7: Save results
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Save qualified keywords analysis
        keywords_output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'qualified_keywords_{timestamp}.xlsx')
        keywords_df = pd.DataFrame(qualifying_keywords)
        keywords_df.to_excel(keywords_output_path, index=False)
        
        # Save qualified ROB data
        rob_output_path = os.path.join(app.config['UPLOAD_FOLDER'], f'qualified_rob_data_{timestamp}.xlsx')
        qualified_rob_data.to_excel(rob_output_path, index=False)
        
        # Step 8: Remove qualifying RIDs from ranking sheet
        log_to_status("Removing qualifying RIDs from ranking sheet...")
        updated_ranking_df = ranking_df[~ranking_df['RID'].isin(qualifying_rids)]
        updated_ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], f'updated_ranking_sheet_{timestamp}.xlsx')
        updated_ranking_df.to_excel(updated_ranking_path, index=False)
        
        # Step 9: Download files to system
        download_folder = os.path.join(os.path.expanduser("~"), "Desktop", "Weekly_Keyword_Analysis")
        os.makedirs(download_folder, exist_ok=True)
        
        # Copy files to download folder
        shutil.copy2(keywords_output_path, os.path.join(download_folder, f'qualified_keywords_{timestamp}.xlsx'))
        shutil.copy2(rob_output_path, os.path.join(download_folder, f'qualified_rob_data_{timestamp}.xlsx'))
        shutil.copy2(updated_ranking_path, os.path.join(download_folder, f'updated_ranking_sheet_{timestamp}.xlsx'))
        
        processing_status['active'] = False
        processing_status['message'] = f"Weekly analysis complete! {len(qualifying_keywords)} keywords qualified, files saved to Desktop/Weekly_Keyword_Analysis"
        
        log_to_status(f"✅ Weekly keyword analysis completed successfully!")
        log_to_status(f"📊 Results:")
        log_to_status(f"   - Qualified keywords: {len(qualifying_keywords)}")
        log_to_status(f"   - Qualified RIDs: {len(qualifying_rids)}")
        log_to_status(f"   - Files saved to: {download_folder}")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Weekly keyword analysis error: {e}")

# Document Processing Routes (Updated for Power Automate workflow)
@app.route('/document_processing', methods=['GET', 'POST'])
def document_processing():
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data - use session data as defaults if available
            article_code = request.form.get('article_code') or request.form.get('open_pr_id') or session.get('open_pr_id', '6HA-2025-M6K439')
            author_name = request.form.get('author_name') or session.get('username', 'Vishwas tiwari')
            author_email = request.form.get('author_email') or session.get('email', 'vishwas@coherentmarketinsights.com')
            company_name = request.form.get('company_name', 'Coherent Market Insights')
            phone_number = request.form.get('phone_number') or session.get('mobile', '1234567890')
            
            # Power Automate output folder path
            custom_folder = request.form.get('custom_folder')
            if custom_folder:
                folder_path = custom_folder
            else:
                today = datetime.today()
                folder_path = rf'C:\Users\akshat\Desktop\RPA\Files\{today.year}\{today.strftime("%m")}\{today.strftime("%d")}'
            
            processing_mode = request.form.get('processing_mode', 'manual')
            
            # Validate paths before processing
            excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
            
            # Check if required files exist
            validation_errors = []
            if not os.path.exists(excel_path):
                validation_errors.append(f"Excel file not found: {excel_path}")
            if not os.path.exists(folder_path):
                validation_errors.append(f"Folder not found: {folder_path}")
            
            if validation_errors:
                for error in validation_errors:
                    flash(error)
                return render_template('document_processing.html', 
                                     session_data={
                                         'username': session.get('username', ''),
                                         'email': session.get('email', ''),
                                         'mobile': session.get('mobile', ''),
                                         'open_pr_id': session.get('open_pr_id', '')
                                     })
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting processing...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread
            if processing_mode == 'auto':
                threading.Thread(target=process_documents_auto_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number)).start()
            else:
                threading.Thread(target=process_documents_manual_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number)).start()
            
            flash('Processing started! Check the status page for updates.')
            return redirect(url_for('processing_status_page'))
            
        except Exception as e:
            flash(f'Error starting processing: {str(e)}')
            logger.error(f"Error in document_processing: {e}")
            return render_template('document_processing.html', 
                                 session_data={
                                     'username': session.get('username', ''),
                                     'email': session.get('email', ''),
                                     'mobile': session.get('mobile', ''),
                                     'open_pr_id': session.get('open_pr_id', '')
                                 })
    
    # Pre-populate form with session data if available
    return render_template('document_processing.html', 
                         session_data={
                             'username': session.get('username', ''),
                             'email': session.get('email', ''),
                             'mobile': session.get('mobile', ''),
                             'open_pr_id': session.get('open_pr_id', '')
                         })

@app.route('/processing_status')
def processing_status_page():
    return render_template('processing_status.html')

@app.route('/api/get_processing_status')
def get_processing_status():
    """API endpoint to get current processing status"""
    global processing_status
    return jsonify(processing_status)

# Document Processing Functions
def convert_doc_to_docx(doc_path, output_path=None):
    """Convert .doc file to .docx format"""
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(doc_path)
        if not output_path:
            output_path = os.path.splitext(doc_path)[0] + ".docx"
        doc.SaveAs(output_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return output_path
    except Exception as e:
        log_to_status(f"Error converting doc to docx: {e}")
        return None

def text_of_press_release(doc_path, start_index=21, end_index=-8):
    # Load the DOCX file
    doc = Document(doc_path)

    # Extract only V4 section paragraphs
    v4_paragraphs = []
    v4_found = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "V4":
            v4_found = True
            continue
        elif text.startswith("Version 5") and v4_found:
            break
        elif v4_found:
            v4_paragraphs.append(para)

    # If V4 not found, fall back to all paragraphs
    if not v4_paragraphs:
        v4_paragraphs = doc.paragraphs

    # Extract text with formatting preservation
    formatted_lines = []
    for para in v4_paragraphs:
        text = para.text.strip()
        if not text or text.replace('_', '').replace('-', '').strip() == "":
            if formatted_lines and formatted_lines[-1] != "":
                formatted_lines.append("")
            continue
        formatted_lines.append(text)

    saved = "\n".join(formatted_lines)
    words = saved.split()

    if len(words) > abs(end_index):
        chunk = " ".join(words[start_index:end_index])

        # Add line breaks before section headers
        section_headers = [
            'Market Size and Overview',
            'Key Takeaways',
            'Segments Covered:',
            'Growth Factors',
            'Market Trends',
            'Actionable Insights',
            'Key Players',
            'FAQs'
        ]
        for header in section_headers:
            chunk = chunk.replace(header, f"\n\n{header}")

        # Remove dashes before content
        chunk = re.sub(r'-{2,}', '', chunk)

        # Add line breaks before bullet points
        chunk = chunk.replace(' - ', '\n- ')

        # Add line breaks before FAQ labels
        chunk = re.sub(r'\s*(FAQ?s?:?)', r'\n\n\1\n\n', chunk)

        # Ensure each numbered FAQ starts on a new line
        # Matches e.g., 1. or 2.
        chunk = re.sub(r'\s*(\d+\.\s)', r'\n\1', chunk)

        chunk = chunk.strip()
        return chunk
    else:
        return "Text not found." 

def run_selenium_automation(article_code, article_title, multiline_text, author_name, 
                          author_email, company_name, phone_number):
    """Run Selenium automation for press release submission"""
    try:
        log_to_status("Starting Selenium automation...")
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        # Uncomment next line for headless mode
        # options.add_argument("--headless")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        pr_agency = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[3]/div/input'))
        )
        pr_agency.clear()
        pr_agency.send_keys(author_name)
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys(company_name)
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        Category_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        Select(Category_element).select_by_visible_text("Arts & Culture")
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        title.clear()
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:
        Mr. Shah
        Coherent Market Insights Pvt. Ltd,
        U.S.: + 12524771362
        U.K.: +442039578553
        AUS: +61-2-4786-0457
        INDIA: +91-848-285-0837
        ✉ Email: sales@coherentmarketinsights.com
        About Us:
        Coherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries.
        """
        about.send_keys(multi)
        
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        address.send_keys("123 Test Street, Test City, Test Country")
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(r"C:\Users\akshat\Desktop\code\image.jpg")
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(5)
        # driver.quit()  # Commented out to keep the browser open for inspection
        log_to_status("Selenium automation completed successfully")
        return True
        
    except Exception as e:
        log_to_status(f"Selenium automation error: {e}")
        try:
            driver.quit()
        except:
            pass
        return False

def process_documents_auto_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number):
    """Process documents automatically with status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting auto processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Auto-processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Convert doc to docx
                processing_status['message'] = f"Converting {market_name} to DOCX..."
                docx_path = convert_doc_to_docx(doc_path)
                
                if not docx_path:
                    log_to_status(f"ERROR: Could not convert {doc_path} to docx")
                    continue
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(docx_path)
                article_title = f"{market_name} Insights"
                
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, 
                                                author_name, author_email, company_name, phone_number)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(10) # Longer delay for auto mode
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Auto-processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Auto processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Auto processing error: {e}")

def process_documents_manual_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number):
    """Process documents with manual intervention and status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting manual processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\akshat\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Convert doc to docx
                processing_status['message'] = f"Converting {market_name} to DOCX..."
                docx_path = convert_doc_to_docx(doc_path)
                
                if not docx_path:
                    log_to_status(f"ERROR: Could not convert {doc_path} to docx")
                    continue
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(docx_path)
                article_title = f"{market_name} Market Insights"
                
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, 
                                                author_name, author_email, company_name, phone_number)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(5)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Manual processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Manual processing error: {e}")

# ROB Processing Routes
@app.route('/rob', methods=['GET', 'POST'])
def rob():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        open_pr_id = request.form.get('open_pr_id')
        mobile = request.form.get('mobile')

        # Validate required fields
        if not all([username, email, open_pr_id, mobile]):
            flash('All fields are required!')
            return redirect(request.url)

        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Excel file is required!')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Only Excel files (.xlsx, .xls) and CSV files are allowed!')
            return redirect(request.url)

        # Use secure_filename to avoid path issues
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Store user data in session for later use
        session['username'] = username
        session['email'] = email
        session['open_pr_id'] = open_pr_id
        session['mobile'] = mobile
        
        return redirect(url_for('process_rob', file_path=input_path,
                                username=username, email=email,
                                open_pr_id=open_pr_id, mobile=mobile))
    return render_template('rob.html')

@app.route('/process_rob')
def process_rob():
    file_path = request.args.get('file_path')
    username = request.args.get('username')
    email = request.args.get('email')
    open_pr_id = request.args.get('open_pr_id')
    mobile = request.args.get('mobile')

    if not file_path or not os.path.exists(file_path):
        flash('Missing or invalid file path')
        return redirect(url_for('rob'))

    try:
        # Read the uploaded file based on extension
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path, engine='openpyxl')

        # List of columns to extract
        columns_to_extract = ['Report ID', 'Report Name', 'Companies covered', 'Market Size Year 2025', 'CAGR', 'Forecast Period', 'Value Projection 2032']

        # Check if the specified columns exist in the DataFrame
        missing_columns = [column for column in columns_to_extract if column not in df.columns]
        if missing_columns:
            flash(f"Missing columns: {', '.join(missing_columns)}")
            return redirect(url_for('rob'))

        # Extract the specified columns and make a copy to avoid SettingWithCopyWarning
        extracted_columns = df[columns_to_extract].copy()

        # Rename the columns
        extracted_columns.rename(columns={
            'Report Name': 'Market Name',
            'Companies covered': 'Key Players',
            'Market Size Year 2025': 'Market Size in 2025'
        }, inplace=True)

        # Set the Forecast Period column to a constant value
        extracted_columns['Forecast Period'] = '2025 to 2032'

        # Add a new column "Market Size Year" with constant value
        extracted_columns['Market Size Year'] = 'Market Size in 2025:'

        # Add a new column "Status" with constant value
        extracted_columns['Status'] = ''

        # Create the new "Market Size" column by combining "Market Size in 2025" and "Value Projection 2032"
        extracted_columns['Market Size'] = extracted_columns['Market Size in 2025'].astype(str) + '; Market Size in 2032: ' + extracted_columns['Value Projection 2032'].astype(str)

        # Create the "Prompt" column with the provided constant text and dynamically generated links
        extracted_columns['Prompt'] = extracted_columns.apply(
            lambda row: f"Furthermore, we have a CTA that needs to be incorporated into the generated blog. Make sure all CTAs are added properly to ensure they are fully synced with the content and, from a lead generation perspective, the placement should be optimal. CTA context: The main link redirects to our main collateral published on our website. The Sample Request URL leads to a page where users can request a sample copy of the report, and the Buy Now URL allows users to directly purchase the report by making a payment. Ensure that the CTAs are placed correctly so they direct the reader to the appropriate webpage linked. The first CTA should be placed after the Market Size and Overview data, the second CTA after the Growth Factors section, and the third CTA after the Actionable Insights. Please do not make any changes to the provided data and links such as do not add brackets, or do not make changes in formatting style, because this blog will be directly published on PR website. CTA Links: First CTA- Explore the Entire Market Report here: https://www.coherentmarketinsights.com/market-insight/{row['Market Name'].replace(' ', '-').lower()}-{row['Report ID']} , 2nd CTA- Request for Sample Copy of the Report here : https://www.coherentmarketinsights.com/insight/request-sample/{row['Report ID']} and 3rd CTA- Get Instant Access! Purchase Research Report and Receive a 25% Discount: https://www.coherentmarketinsights.com/insight/buy-now/{row['Report ID']}",
            axis=1
        )

        # Combine Key Players with Prompt text (with single space)
        extracted_columns['Key Players'] = extracted_columns['Key Players'].astype(str) + ' ' + extracted_columns['Prompt']
        
        # Prepare the final output DataFrame (removed 'Prompt' since it's now combined with Key Players)
        output_df = extracted_columns[['Market Name', 'Key Players', 'Market Size Year', 'Market Size', 'CAGR', 'Forecast Period', 'Status', 'Report ID']]

        # Desired order of columns
        desired_order = ['Market Name', 'Forecast Period', 'Market Size Year', 'Market Size', 'CAGR', 'Key Players', 'Status', 'Report ID']

        # Reorder columns according to the desired order
        output_df = output_df[desired_order]

        # Save the final DataFrame to an Excel file in the RPA folder on Desktop
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        # Save the final DataFrame to an Excel file in the RPA folder
        output_path = os.path.join(rpa_folder, "ROB.xlsx")
        output_df.to_excel(output_path, index=False)
        
        flash(f"ROB processing complete. File saved to: {output_path}")
        flash(f"Total rows processed: {len(output_df)}")
        
        return redirect(url_for('wait_power_automate'))

    except Exception as e:
        flash(f"Error processing ROB file: {str(e)}")
        return redirect(url_for('rob'))

@app.route('/wait_power_automate')
def wait_power_automate():
    """Show a waiting page for Power Automate Desktop step."""
    return render_template('wait_power_automate.html')

@app.route('/api/trigger_power_automate', methods=['POST'])
def trigger_power_automate_flow():
    """
    Triggers a Power Automate Desktop flow by launching the PAD executable and running the specified flow.
    """
    # Path to the PAD.ConsoleHost.exe (check if this path is correct on your system)
    pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
    flow_name = "Paid PR - Files Downloader"
    # Verify if the PAD executable exists
    if not os.path.exists(pad_exe_path):
        print("Power Automate Desktop executable not found!")
        return jsonify({'status': 'error', 'message': 'Power Automate Desktop not found'})
    
    # Construct the command to trigger the flow
    command = f'"{pad_exe_path}" -flow "{flow_name}"'
    
    # Execute the command using subprocess
    try:
        result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
        print(f"Flow triggered successfully. Output: {result.stdout}")

        # Wait for PAD to load
        time.sleep(5)
        
        # Use PyAutoGUI to click the flow
        flow_button_coordinates = (463, 395)
        print(f"Clicking at {flow_button_coordinates}")
        pyautogui.click(flow_button_coordinates)
        print("Flow triggered successfully.")

    except subprocess.CalledProcessError as e:
        print(f"Error triggering flow: {e.stderr}")
        return jsonify({'status': 'error', 'message': f'Flow trigger failed: {e.stderr}'})
    
    return jsonify({'status': 'success', 'message': 'Power Automate process completed.'})

if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True, host='0.0.0.0', port=5000)
